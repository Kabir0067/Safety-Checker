from typing import Optional, Dict, List, Any,Tuple
from asyncio import Semaphore
from docx import Document
import concurrent.futures
from pathlib import Path
from PIL import Image, ImageEnhance
import pandas as pd
import numpy as np
import mimetypes
import aiofiles
import easyocr
import logging
import asyncio
import shutil 
import cv2
import os
import re
import docx   
import pytesseract
import fitz





class ProfessionalOCRProcessor:
    LOG_DIR = "logs"
    LOG_FILE = os.path.join(LOG_DIR, "professional_ocr_processor.log")
    
    def __init__(
        self, 
        languages: List[str] = ['eng', 'rus'],
        min_confidence: float = 0.6,
        gpu: bool = False,
        model_storage_directory: str = None,
        download_enabled: bool = True
    ):
        self.languages = "+".join(languages)
        self.min_confidence = min_confidence
        self.supported_formats = {'.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.webp', '.jfif'}
        
        self.tesseract_config = f'--oem 3 --psm 3'
        
        self._setup_logging()
        
        os.makedirs(self.LOG_DIR, exist_ok=True)
        os.makedirs("tmp/processed", exist_ok=True)
        os.makedirs("tmp/debug", exist_ok=True)
        
        self.logger.info(f"ProfessionalOCRProcessor инициализирован с языками: {self.languages}")
        try:
            pytesseract.get_tesseract_version()
            self.logger.info(f"Tesseract version: {pytesseract.get_tesseract_version()}")
        except pytesseract.TesseractNotFoundError:
            self.logger.error("❌ Tesseract насб нашудааст ё дар PATH нест. Лутфан, пеш аз идома додан Tesseract-ро насб кунед.")

    def _setup_logging(self):
        self.logger = logging.getLogger("ProfessionalOCRProcessor")
        self.logger.setLevel(logging.INFO)
        formatter = logging.Formatter(
            '%(asctime)s - %(name)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s'
        )
        file_handler = logging.FileHandler(self.LOG_FILE, encoding='utf-8')
        file_handler.setLevel(logging.INFO)
        file_handler.setFormatter(formatter)
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        console_handler.setFormatter(formatter)
        if self.logger.handlers:
            self.logger.handlers.clear()
        self.logger.addHandler(file_handler)
        self.logger.addHandler(console_handler)

    def cleanup_tmp(self):
        for tmp_dir in ["tmp/processed", "tmp/debug"]:
            if os.path.exists(tmp_dir):
                try:
                    shutil.rmtree(tmp_dir)
                    self.logger.info(f"Папка пок карда шуд: {tmp_dir}")
                except Exception as e:
                    self.logger.warning(f"Хатогӣ ҳангоми пок кардани {tmp_dir}: {e}")

    def __del__(self):
        try:
            self.cleanup_tmp()
        except:
            pass

    def validate_image(self, image_path: str) -> Tuple[bool, str]:
        try:
            if not os.path.exists(image_path):
                return False, f"Файл не существует: {image_path}"
                
            file_ext = Path(image_path).suffix.lower()
            if file_ext not in self.supported_formats:
                return False, f"Неподдерживаемый формат: {file_ext}"
            
            file_size = os.path.getsize(image_path) / (1024 * 1024) 
            if file_size > 100: 
                return False, f"Файл слишком большой: {file_size:.1f}MB"
            
            with Image.open(image_path) as img:
                img.verify()
                if img.size[0] < 10 or img.size[1] < 10:
                    return False, "Изображение слишком маленькое"
                if img.size[0] > 10000 or img.size[1] > 10000:
                    return False, "Изображение слишком большое"
            
            self.logger.info(f"Файл валиден: {image_path} ({file_size:.2f}MB)")
            return True, "OK"
            
        except Exception as e:
            error_msg = f"Ошибка валидации изображения {image_path}: {str(e)}"
            self.logger.error(error_msg)
            return False, error_msg

    def load_image(self, path: str) -> Optional[np.ndarray]:
        try:
            with open(path, 'rb') as f:
                data = np.frombuffer(f.read(), dtype=np.uint8)
            img = cv2.imdecode(data, cv2.IMREAD_COLOR)
            if img is None:
                raise ValueError("Не удалось декодировать изображение")
            return img
        except Exception as e:
            self.logger.error(f"Ошибка загрузки изображения {path}: {e}")
            return None

    def save_debug_image(self, image: np.ndarray, prefix: str, filename: str):
        try:
            debug_path = f"tmp/debug/{prefix}_{Path(filename).stem}.png"
            cv2.imwrite(debug_path, image)
            return debug_path
        except Exception as e:
            self.logger.warning(f"Не удалось сохранить отладочное изображение: {e}")
            return None

    def advanced_preprocessing(self, image_path: str) -> Optional[np.ndarray]:
        try:
            is_valid, message = self.validate_image(image_path)
            if not is_valid:
                self.logger.error(message)
                return None

            img = self.load_image(image_path)
            if img is None:
                return None

            original_size = img.shape[:2]
            self.logger.info(f"Оригинальный размер: {original_size[1]}x{original_size[0]}")

            if len(img.shape) == 3:
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            else:
                gray = img

            max_dimension = max(original_size)
            scale_factor = 1.0
            if max_dimension < 1000:
                scale_factor = 2.0
            elif max_dimension < 3000:
                scale_factor = 1.5
            
            if scale_factor != 1.0:
                new_size = (int(original_size[1] * scale_factor), int(original_size[0] * scale_factor))
                gray = cv2.resize(gray, new_size, interpolation=cv2.INTER_LANCZOS4) 
                self.logger.info(f"Масштабированный размер: {new_size[0]}x{new_size[1]}")

            clahe = cv2.createCLAHE(clipLimit=4.0, tileGridSize=(12, 12))
            enhanced = clahe.apply(gray)

            gaussian = cv2.GaussianBlur(enhanced, (0, 0), 5.0)
            sharpened = cv2.addWeighted(enhanced, 2.0, gaussian, -1.0, 0)

            block_size = 21
            if min(sharpened.shape) < 500:
                 block_size = 11
            
            block_size = block_size if block_size % 2 != 0 else block_size + 1 
            binary = cv2.adaptiveThreshold(
                sharpened, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY, block_size, 2
            )

            kernel_open = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 1))
            kernel_close = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
            
            cleaned = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel_open)
            cleaned = cv2.morphologyEx(cleaned, cv2.MORPH_CLOSE, kernel_close)
            cleaned = cv2.medianBlur(cleaned, 3)

            self.save_debug_image(cleaned, "advanced", Path(image_path).name)

            return cleaned

        except Exception as e:
            self.logger.exception(f"Ошибка в advanced_preprocessing: {e}")
            return None

    def medium_preprocessing(self, image_path: str) -> Optional[np.ndarray]:
        return self.advanced_preprocessing(image_path)

    def simple_preprocessing(self, image_path: str) -> Optional[np.ndarray]:
        try:
            is_valid, message = self.validate_image(image_path)
            if not is_valid:
                return None

            img = self.load_image(image_path)
            if img is None:
                return None

            if len(img.shape) == 3:
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            else:
                gray = img

            height, width = gray.shape
            if max(height, width) < 1000:
                new_size = (int(width * 1.5), int(height * 1.5))
                gray = cv2.resize(gray, new_size, interpolation=cv2.INTER_LANCZOS4)

            self.save_debug_image(gray, "simple", Path(image_path).name)

            return gray

        except Exception as e:
            self.logger.exception(f"Ошибка в simple_preprocessing: {e}")
            return None

    def create_image_variants(self, base_image: np.ndarray) -> List[Tuple[str, np.ndarray]]:
        variants = [("original", base_image)]
        
        try:
            inverted = cv2.bitwise_not(base_image)
            variants.append(("inverted", inverted))
            
            if base_image.dtype == np.uint8:
                high_contrast = cv2.convertScaleAbs(base_image, alpha=2.0, beta=0) 
                variants.append(("high_contrast", high_contrast))
                
        except Exception as e:
            self.logger.warning(f"Не удалось создать некоторые варианты изображения: {e}")
            
        return variants

    def extract_specific_data(self, text: str) -> Dict[str, List[str]]:
        data = {
            'phone_numbers': [],
            'emails': [],
            'domains': []
        }
        
        phone_pattern = re.compile(
            r'(\+?\d{1,3}[-. ]?)?(\(?\d{3}\)?[-. ]?)(\d{3}[-. ]?)(\d{4})', 
            re.IGNORECASE
        )
        
        email_pattern = re.compile(
            r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
            re.IGNORECASE
        )
        
        domain_pattern = re.compile(
            r'(?:https?://|www\.)?([a-zA-Z0-9-]+\.(?:com|org|net|info|biz|co|tj|ru|uz|kz|[\w]{2,}))',
            re.IGNORECASE
        )

        for match in phone_pattern.finditer(text):
            full_match = match.group(0).strip()
            digits = re.sub(r'[^\d]', '', full_match)
            if len(digits) >= 7:
                 data['phone_numbers'].append(full_match)

        data['emails'] = list(set(email_pattern.findall(text)))
        
        all_urls_and_domains = domain_pattern.findall(text)
        data['domains'] = list(set([d.lower() for d in all_urls_and_domains if '.' in d]))
        
        return data

    def postprocess_text(self, text: str) -> str:
        if not text or not text.strip():
            return ""

        # 1. Тозакунии фазо
        text = re.sub(r'\s+', ' ', text).strip()
        
        # 2. Ислоҳоти маъмулии OCR (expanded for common confusions: O/0, I/1/l, S/5, B/8, etc.)
        corrections = [
            (r'\bI\b', '1'),  # I to 1 in numbers
            (r'\|', 'l'), 
            (r'\[', 'l'), 
            (r'\]', 'l'), 
            (r'©', 'c'), 
            (r'®', 'r'), 
            (r'O(?=\d)', '0'),  # O to 0 before digits
            (r'(\d)O', r'\g<1>0'),  # O to 0 after digits
            (r'S(?=\d)', '5'),  # S to 5 in numbers
            (r'(\d)S', r'\g<1>5'),  # S to 5 after digits
            (r'B(?=\d)', '8'),  # B to 8 in numbers
            (r'(\d)B', r'\g<1>8'),  # B to 8 after digits
            (r'l(?=\d)', '1'),  # l to 1 in numbers
            (r'(\d)l', r'\g<1>1'),  # l to 1 after digits
            (r'Z(?=\d)', '2'),  # Z to 2
            (r'(\d)Z', r'\g<1>2'),
            (r'G(?=\d)', '6'),  # G to 6
            (r'(\d)G', r'\g<1>6'),
            (r'\bO\b', '0'),    # Standalone O to 0
            (r',O', ',0'),      # Comma O to 0 in numbers like 52,O00
            (r'\b[sS]\b', 'S'), # Fix case for single letters
            (r'l', 'I'),        # l to I in some contexts
        ]
        
        for pattern, replacement in corrections:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)  # Added IGNORECASE

        # Additional fixes for symbols and punctuation
        text = re.sub(r'\~', '-', text)  # Tilde to dash
        text = re.sub(r'—', '-', text)   # Em dash to hyphen
        text = re.sub(r'’', "'", text)   # Smart quotes to straight
        text = re.sub(r'“|”', '"', text) 

        # 3. Пайваст кардани хатҳои тақсимшуда
        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
        text = re.sub(r'(\w+)-\s+(\w+)', r'\1\2', text)
        text = re.sub(r'(\w+)-\s*\n', r'\1', text)

        # 4. Капитализатсия ва тозакунии охирин
        lines = []
        for line in text.split('\n'):
            line = line.strip()
            if len(line) > 0:
                if len(line) > 1 and line[0].islower():
                    line = line[0].upper() + line[1:]
                lines.append(line)
                
        cleaned_text = '\n'.join(lines)
        
        cleaned_text = re.sub(r' +', ' ', cleaned_text)
        cleaned_text = re.sub(r'\n\s*\n', '\n\n', cleaned_text)

        # 5. Simple spell correction for common words (no external lib, manual dict)
        common_corrections = {
            'HsBC': 'HSBC',
            'sW1A': 'SW1A',
            '2O25l': '2025',
            # Add more if needed based on domain
        }
        for wrong, right in common_corrections.items():
            cleaned_text = re.sub(re.escape(wrong), right, cleaned_text)

        return cleaned_text

    def calculate_text_quality(self, text: str, avg_confidence: float) -> Dict[str, float]:
        if not text:
            return {'overall': 0.0, 'structure': 0.0, 'readability': 0.0, 'confidence': 0.0}

        words = text.split()
        word_count = len(words)
        lines = text.split('\n')
        
        overall_score = avg_confidence * 0.7 
        
        if word_count > 50:
            overall_score += 0.2
            
        specific_data = self.extract_specific_data(text)
        if specific_data['phone_numbers'] or specific_data['emails'] or specific_data['domains']:
             overall_score += 0.1 
             
        overall_score = min(overall_score, 1.0)
        
        return {
            'overall': overall_score,
            'structure': min(len(lines) / 10, 1.0),
            'readability': min(word_count / 100, 1.0),
            'confidence': avg_confidence
        }
    
    def perform_tesseract_ocr(self, image: np.ndarray) -> Tuple[str, float]:
        try:
            data = pytesseract.image_to_data(
                image, 
                lang=self.languages, 
                config=self.tesseract_config, 
                output_type=pytesseract.Output.DATAFRAME
            )
            
            data = data[data.conf != -1]
            data = data.dropna(subset=['text'])
            
            reliable_data = data[data.conf / 100 >= self.min_confidence]
            
            if reliable_data.empty:
                full_text = pytesseract.image_to_string(image, lang=self.languages, config=self.tesseract_config)
                avg_confidence = data['conf'].mean() / 100 if not data.empty else 0.0
                return full_text, avg_confidence

            text = ' '.join(reliable_data['text'].astype(str))
            
            avg_confidence = reliable_data['conf'].mean() / 100
            
            return text, avg_confidence
            
        except pytesseract.TesseractNotFoundError:
            self.logger.error("❌ Tesseract насб нашудааст ё дар PATH нест.")
            return "", 0.0
        except Exception as e:
            self.logger.error(f"Ошибка Tesseract OCR: {e}")
            return "", 0.0

    def extract_text_from_results(self, text: str, confidence: float) -> Tuple[str, float]:
        return text, confidence 

    def perform_ocr_with_fallback(self, file_path: str) -> Dict[str, Any]:
        self.logger.info(f"Начало OCR обработки: {file_path}")
        
        all_results = []
        preprocessing_methods = [
            ("simple", self.simple_preprocessing),
            ("advanced", self.advanced_preprocessing),  # Removed medium as it's duplicate
        ]

        try:
            for method_name, preprocess_func in preprocessing_methods:
                try:
                    self.logger.info(f"Пробуем метод препроцессинга: {method_name}")
                    processed_image = preprocess_func(file_path)
                    
                    if processed_image is None:
                        continue
                        
                    image_variants = self.create_image_variants(processed_image)
                    
                    for variant_name, variant_image in image_variants:
                        try:
                            raw_text, avg_confidence = self.perform_tesseract_ocr(variant_image)
                            
                            if raw_text and avg_confidence >= self.min_confidence:
                                cleaned_text = self.postprocess_text(raw_text)
                                quality_scores = self.calculate_text_quality(cleaned_text, avg_confidence)
                                specific_data = self.extract_specific_data(cleaned_text)
                                
                                result_data = {
                                    'text': cleaned_text,
                                    'raw_text': raw_text,
                                    'confidence': avg_confidence,
                                    'quality_scores': quality_scores,
                                    'preprocessing_method': method_name,
                                    'image_variant': variant_name,
                                    'word_count': len(cleaned_text.split()),
                                    'character_count': len(cleaned_text),
                                    'specific_data': specific_data,
                                }
                                
                                all_results.append(result_data)
                                self.logger.info(
                                    f"Найден текст ({method_name}, {variant_name}): "
                                    f" ({avg_confidence:.3f})"
                                )
                                
                        except Exception as e:
                            self.logger.warning(f"Ошибка при обработке варианта {variant_name}: {e}")
                            continue
                            
                except Exception as e:
                    self.logger.exception(f"Ошибка в методе {method_name}: {e}")
                    continue

            if not all_results:
                self.logger.info("Пробуем фолбэк: обработка без препроцессинга")
                try:
                    original_image = self.load_image(file_path)
                    if original_image is not None:
                        raw_text, avg_confidence = self.perform_tesseract_ocr(original_image)
                        if raw_text and avg_confidence > 0.1:
                            cleaned_text = self.postprocess_text(raw_text)
                            quality_scores = self.calculate_text_quality(cleaned_text, avg_confidence)
                            specific_data = self.extract_specific_data(cleaned_text)
                            
                            all_results.append({
                                'text': cleaned_text,
                                'raw_text': raw_text,
                                'confidence': avg_confidence,
                                'quality_scores': quality_scores,
                                'preprocessing_method': 'fallback',
                                'image_variant': 'original',
                                'word_count': len(cleaned_text.split()),
                                'character_count': len(cleaned_text),
                                'specific_data': specific_data
                            })
                except Exception as e:
                    self.logger.error(f"Фолбэк метод не сработал: {e}")

            if all_results:
                all_results.sort(key=lambda x: x['quality_scores']['overall'] * x['confidence'], reverse=True)
                best_result = all_results[0]
                
                self.logger.info(
                    f"Лучший результат: метод={best_result['preprocessing_method']}, "
                    f"вариант={best_result['image_variant']}, "
                    f"уверенность={best_result['confidence']:.3f}, "
                    f"слов={best_result['word_count']}, "
                    f"оценка качества={best_result['quality_scores']['overall']:.3f}"
                )
                
                result = {
                    'status': 'success',
                    'text': best_result['text'],
                    'raw_text': best_result.get('raw_text', ''),
                    'confidence': best_result['confidence'],
                    'quality_scores': best_result['quality_scores'],
                    'preprocessing_method': best_result['preprocessing_method'],
                    'image_variant': best_result['image_variant'],
                    'word_count': best_result['word_count'],
                    'character_count': best_result['character_count'],
                    'all_attempts': len(all_results),
                    'timestamp': pd.Timestamp.now().isoformat(),
                    'specific_data': best_result['specific_data']
                }
            else:
                self.logger.error("Не удалось извлечь текст ни одним методом")
                result = {
                    'status': 'error',
                    'text': 'No reliable text found with any method',
                    'confidence': 0.0,
                    'quality_scores': {'overall': 0.0, 'structure': 0.0, 'readability': 0.0},
                    'preprocessing_method': 'none',
                    'word_count': 0,
                    'character_count': 0,
                    'all_attempts': 0,
                    'timestamp': pd.Timestamp.now().isoformat(),
                    'specific_data': {'phone_numbers': [], 'emails': [], 'domains': []}
                }

            return result

        finally:
            self.cleanup_tmp()

    async def perform_ocr_async(self, file_path: str) -> Dict[str, Any]:
        try:
            self.logger.info(f"Асинхронная обработка: {file_path}")
            loop = asyncio.get_event_loop()
            
            with concurrent.futures.ThreadPoolExecutor() as pool:
                result = await asyncio.wait_for(
                    loop.run_in_executor(pool, self.perform_ocr_with_fallback, file_path),
                    timeout=180.0  
                )
            return result
            
        except asyncio.TimeoutError:
            self.logger.error(f"Таймаут обработки: {file_path}")
            return {
                'status': 'error',
                'text': 'OCR processing timeout',
                'confidence': 0.0,
                'quality_scores': {'overall': 0.0, 'structure': 0.0, 'readability': 0.0},
                'specific_data': {'phone_numbers': [], 'emails': [], 'domains': []}
            }
        except Exception as e:
            self.logger.exception(f"Ошибка асинхронной обработки {file_path}: {e}")
            return {
                'status': 'error',
                'text': f'OCR processing failed: {str(e)}',
                'confidence': 0.0,
                'quality_scores': {'overall': 0.0, 'structure': 0.0, 'readability': 0.0},
                'specific_data': {'phone_numbers': [], 'emails': [], 'domains': []}
            }
        finally:
            try:
                loop = asyncio.get_running_loop()
            except RuntimeError:
                loop = asyncio.new_event_loop()
            loop.run_in_executor(None, self.cleanup_tmp)

    def batch_process(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        results = []
        total_files = len(file_paths)
        
        try:
            for i, file_path in enumerate(file_paths, 1):
                try:
                    self.logger.info(f"Обработка файла {i}/{total_files}: {file_path}")
                    result = self.perform_ocr_with_fallback(file_path)
                    results.append({
                        'file_path': file_path,
                        'result': result,
                        'processing_order': i
                    })
                    
                except Exception as e:
                    self.logger.exception(f"Ошибка обработки {file_path}: {e}")
                    results.append({
                        'file_path': file_path,
                        'result': {
                            'status': 'error',
                            'text': f'Processing failed: {str(e)}'
                        },
                        'processing_order': i
                    })
                    
            return results
        finally:
            self.cleanup_tmp()

    async def batch_process_async(self, file_paths: List[str], max_concurrent: int = 3) -> List[Dict[str, Any]]:
        semaphore = Semaphore(max_concurrent)
        
        async def process_with_semaphore(file_path: str) -> Dict[str, Any]:
            async with semaphore:
                result = await self.perform_ocr_async(file_path)
                return {'file_path': file_path, 'result': result}
        
        tasks = [process_with_semaphore(file_path) for file_path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        formatted_results = []
        try:
            for i, (file_path, result) in enumerate(zip(file_paths, results), 1):
                if isinstance(result, Exception):
                    self.logger.exception(f"Исключение при обработке {file_path}: {result}")
                    formatted_results.append({
                        'file_path': file_path,
                        'result': {
                            'status': 'error',
                            'text': f'Exception: {str(result)}'
                        },
                        'processing_order': i
                    })
                else:
                    formatted_results.append({
                        'file_path': file_path,
                        'result': result['result'],
                        'processing_order': i
                    })
            
            return formatted_results
        finally:
            try:
                loop = asyncio.get_running_loop()
            except RuntimeError:
                loop = asyncio.new_event_loop()
            loop.run_in_executor(None, self.cleanup_tmp)

    def export_results(self, results: List[Dict[str, Any]], output_format: str = 'json') -> Any:
        try:
            if output_format == 'json':
                return results
            elif output_format == 'dataframe':
                df_data = []
                for result in results:
                    if result['result']['status'] == 'success':
                        df_data.append({
                            'file_path': result['file_path'],
                            'text': result['result']['text'],
                            'confidence': result['result']['confidence'],
                            'word_count': result['result']['word_count'],
                            'preprocessing_method': result['result']['preprocessing_method'],
                            'phone_numbers': ', '.join(result['result']['specific_data']['phone_numbers']),
                            'emails': ', '.join(result['result']['specific_data']['emails']),
                            'domains': ', '.join(result['result']['specific_data']['domains'])
                        })
                return pd.DataFrame(df_data)
            elif output_format == 'text':
                text_output = []
                for result in results:
                    text_output.append(f"File: {result['file_path']}")
                    text_output.append(f"Status: {result['result']['status']}")
                    if result['result']['status'] == 'success':
                        text_output.append(f"Text: {result['result']['text'][:100]}...")
                        text_output.append(f"Phones: {', '.join(result['result']['specific_data']['phone_numbers'])}")
                        text_output.append(f"Emails: {', '.join(result['result']['specific_data']['emails'])}")
                        text_output.append(f"Domains: {', '.join(result['result']['specific_data']['domains'])}")
                    text_output.append("---")
                return '\n'.join(text_output)
            else:
                raise ValueError(f"Unsupported format: {output_format}")
                
        except Exception as e:
            self.logger.error(f"Ошибка экспорта результатов: {e}")
            return None






class FileConvertToText:
    FILES_DIR = "files"
    MAX_SIZE_BYTES = 10 * 1024 * 1024 
    SUPPORTED_IMAGE_EXTENSIONS = ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.webp', '.jfif']
    SUPPORTED_FORMATS = {
        'word': ['.doc', '.docx'],
        'pdf': ['.pdf'],
        'spreadsheet': ['.csv', '.xls', '.xlsx'],
        'text': ['.txt', '.text'],
        'image': SUPPORTED_IMAGE_EXTENSIONS
    }
    LOG_DIR = "logs"
    LOG_FILE = os.path.join(LOG_DIR, "file_convert_to_text_errors.log")

    def __init__(self):
        os.makedirs(self.FILES_DIR, exist_ok=True)
        os.makedirs(self.LOG_DIR, exist_ok=True)
        self.logger = logging.getLogger("FileConvertToText")
        self.logger.setLevel(logging.ERROR)
        handler = logging.FileHandler(self.LOG_FILE)
        handler.setLevel(logging.ERROR)
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        handler.setFormatter(formatter)
        self.logger.addHandler(handler)
        self.ocr_processor = ProfessionalOCRProcessor(
            languages=['eng', 'rus'],
            min_confidence=0.5
            )

    async def get_file_format(self, file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        if not path.exists() or not path.is_file():
            return {"error": "File does not exist", "status": "error"}
        
        file_size = path.stat().st_size
        if file_size > self.MAX_SIZE_BYTES:
            return {"error": "File is too large (max 10 MB)", "status": "error"}
        
        mime_type, _ = mimetypes.guess_type(file_path)
        file_extension = str(path.suffix.lower())
        
        return {
            "status": "success",
            "extension": file_extension,
            "mime_type": mime_type or "unknown",
            "size_bytes": file_size,
            "size_human": f"{file_size / (1024 * 1024):.2f} MB"
        }

  
    async def read_word(self, file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        if not path.exists() or not path.is_file():
            return {"status": "error", "text": "File does not exist", "metadata": {}}

        if path.stat().st_size > self.MAX_SIZE_BYTES:
            return {"status": "error", "text": "File too large (max 10 MB)", "metadata": {}}

    async def extract_docx_async(self, file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        if not path.exists() or not path.is_file():
            return {"status": "error", "text": f"File {file_path} not found.", "metadata": {}}
        if path.stat().st_size > self.MAX_SIZE_BYTES:
            return {"status": "error", "text": f"{path.name}: file too large.", "metadata": {}}

        def extract_docx():
            try:
                doc = docx.Document(str(path))
                paragraphs = [p.text if p.text.strip() else "" for p in doc.paragraphs]
                text = "\n".join(paragraphs)
                metadata = {
                    "paragraph_count": len(paragraphs),
                    "word_count": len(text.split()),
                    "source": "python-docx"
                }
                return {"status": "success", "text": text, "metadata": metadata}
            except Exception as e:
                return {"status": "error", "text": f"python-docx failed: {str(e)}", "metadata": {}}

        try:
            return await asyncio.to_thread(extract_docx)
        except Exception as e:
            return {"status": "error", "text": f"Error: {str(e)}", "metadata": {}}

    async def pdf_to_text_async(self, file_path: str, output_dir: Optional[str] = None) -> Dict[str, Any]:
        path = Path(file_path)
        if not path.exists() or not path.is_file():
            return {"status": "error", "text": f"File {file_path} not found.", "metadata": {}}
        if path.stat().st_size > self.MAX_SIZE_BYTES:
            return {"status": "error", "text": f"{path.name}: file too large.", "metadata": {}}

        def convert_pdf():
            try:
                doc = fitz.open(str(path))
                text = "".join([page.get_text() for page in doc])
                return {"status": "success", "text": text, "metadata": {"page_count": len(doc)}}
            except Exception as e:
                return {"status": "error", "text": str(e), "metadata": {}}

        return await asyncio.to_thread(convert_pdf)

    async def docx_to_text_async(self, file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        if not path.exists() or not path.is_file():
            return {"status": "error", "text": f"File {file_path} not found.", "metadata": {}}
        if path.stat().st_size > self.MAX_SIZE_BYTES:
            return {"status": "error", "text": f"{path.name}: file too large.", "metadata": {}}

        def convert_docx():
            try:
                doc = docx.Document(str(path))
                text = "\n".join([p.text for p in doc.paragraphs])
                return {"status": "success", "text": text, "metadata": {"paragraphs": len(doc.paragraphs)}}
            except Exception as e:
                return {"status": "error", "text": str(e), "metadata": {}}

        return await asyncio.to_thread(convert_docx)
    async def read_csv_or_excel(self, file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        if not path.exists() or not path.is_file():
            return {"status": "error", "text": "File does not exist", "metadata": {}}
            
        if path.stat().st_size > self.MAX_SIZE_BYTES:
            return {"status": "error", "text": "File too large (max 10 MB)", "metadata": {}}

        def extract_table():
            try:
                if path.suffix.lower() == '.csv':
                    df = pd.read_csv(path, keep_default_na=False)
                elif path.suffix.lower() in ['.xls', '.xlsx']:
                    df = pd.read_excel(path, keep_default_na=False)
                else:
                    raise ValueError("Unsupported")

                lines = []
                header = " | ".join(df.columns.astype(str))
                lines.append(header)
                lines.append("-|-".join(["-" * len(col) for col in df.columns.astype(str)]))
                for _, row in df.iterrows():
                    lines.append(" | ".join(str(cell) for cell in row))
                
                text = "\n".join(lines)
                metadata = {
                    "row_count": len(df),
                    "column_count": len(df.columns),
                    "columns": list(df.columns)
                }
                return {"status": "success", "text": text, "metadata": metadata}
            except Exception as e:
                raise Exception(f"Error: {str(e)}")

        try:
            return await asyncio.to_thread(extract_table)
        except Exception as e:
            self.logger.exception(e)
            return {"status": "error", "text": str(e), "metadata": {}}

    
    async def read_text_file(self, file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        if not path.exists() or not path.is_file():
            return {"status": "error", "text": "File does not exist", "metadata": {}}
        
        encodings = ['utf-8', 'cp1251', 'latin-1']
        
        for encoding in encodings:
            try:
                async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                    text = await f.read()
                    text = text.replace('\r\n', '\n').replace('\r', '\n')
                    metadata = {
                        "line_count": text.count('\n') + 1,
                        "word_count": len(re.findall(r'\b\w+\b', text)),
                        "encoding_used": encoding
                    }
                    return {"status": "success", "text": text, "metadata": metadata}
            except UnicodeDecodeError:
                continue
            except Exception as e:
                self.logger.exception(e)
                break
        
        return {"status": "error", "text": "Failed to decode file", "metadata": {}}


    async def read_image_to_text(self, file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        if not path.exists(): 
            return {"status": "error", "text": f"File not found: {file_path}", "metadata": {}}
        
        if path.stat().st_size > self.MAX_SIZE_BYTES: 
            return {"status": "error", "text": f"{path.name}: too large", "metadata": {}}
        
        if path.suffix.lower() not in self.SUPPORTED_IMAGE_EXTENSIONS:
            return {"status": "error", "text": f"Unsupported image type: {path.suffix}", "metadata": {}}
        
        if self.ocr_processor is None:
            return {"status": "error", "text": "OCR Processor not initialized", "metadata": {}}
        
        print(f"Starting OCR for: {path.name}")
        
        result = await self.ocr_processor.perform_ocr_async(str(path))
        
        if isinstance(result, dict) and result.get('status') == 'success':
            text = result.get('text', '')
            metadata = result.get('metadata', {})
            return {"status": "success", "text": text, "metadata": metadata}
        
        if isinstance(result, str):
            return {"status": "success", "text": result, "metadata": {"source": "ocr_string"}}
        
        return {"status": "error", "text": "OCR failed", "metadata": {}}


    async def convert_to_text(self, file_path: str) -> Dict[str, Any]:
        file_info = await self.get_file_format(file_path)
        if file_info.get("status") == "error":
            return {"status": "error", "text": file_info["error"], "metadata": {}}

        ext = file_info.get("extension", "").lower()
        
        if ext in self.SUPPORTED_FORMATS['word']: 
            return await self.read_word(file_path)
        elif ext in self.SUPPORTED_FORMATS['pdf']: 
            return await self.pdf_to_text_async(file_path)
        elif ext in self.SUPPORTED_FORMATS['spreadsheet']: 
            return await self.read_csv_or_excel(file_path)
        elif ext in self.SUPPORTED_FORMATS['text']:
            return await self.read_text_file(file_path)
        elif ext in self.SUPPORTED_FORMATS['image']:
            return await self.read_image_to_text(file_path)
        else:
            return {"status": "error", "text": f"Unsupported format: {ext}", "metadata": {}}


    async def process_multiple_files(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        tasks = [self.convert_to_text(fp) for fp in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        processed = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                self.logger.exception(result)
                processed.append({"status": "error", "file": file_paths[i], "text": str(result), "metadata": {}})
            else:
                processed.append({**result, "file": file_paths[i]})
        return processed
    





    